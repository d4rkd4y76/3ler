#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""3. sınıf Kelime Dağarcığı — ortak build yardımcıları."""
from __future__ import annotations

import json
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

VISUAL_MARKUP_RE = re.compile(
    r"\[\[(?:clock|clockarc|para|terazi|shape|solid|bunny):[^\]]+\]\]",
    re.IGNORECASE,
)


def premise_plain_text(premise: str) -> str:
    s = VISUAL_MARKUP_RE.sub("", premise or "")
    return re.sub(r"\s+", " ", s).strip()


def has_visual_markup(premise: str) -> bool:
    return bool(VISUAL_MARKUP_RE.search(premise or ""))


def premise_text_prefix(premise: str) -> str | None:
    """Öncül tipine göre soru kökü öneki."""
    if not premise:
        return None
    plain = premise_plain_text(premise)
    if parse_roman_items(premise):
        return None
    if has_visual_markup(premise) and len(plain) < 20:
        return "Yukarıdakine göre"
    if len(plain) > 15 or (len(premise) > 15 and not has_visual_markup(premise)):
        return "Yukarıdaki metne göre"
    if has_visual_markup(premise):
        return "Yukarıdakine göre"
    return None


def parse_roman_items(text: str) -> list[dict]:
    items = []
    for line in (text or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^((?:I{1,3}|IV|VI{0,3}|IX|X)\.)\s*(.+)$", line)
        if m:
            items.append({"label": m.group(1), "text": m.group(2).strip()})
    return items


def normalize_question(q: dict) -> dict:
    """Öncül/soru ayrımını düzelt — cümle metni öncüle, soru kökü ayrı kalır."""
    import copy

    q = copy.copy(q)
    text = (q.get("text") or "").strip()
    premise = (q.get("premise") or "").strip() or None

    m = re.match(
        r"^(.*?)\s*Cümle:\s*(.+)$",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if m:
        stem = m.group(1).strip()
        sentence = m.group(2).strip()
        if "Yukarı" not in stem and "yukarı" not in stem.lower():
            stem = stem.replace(
                "Boş bırakılan yere",
                "Yukarıdaki cümlede boş bırakılan yere",
            )
            if "Yukarı" not in stem:
                stem = "Yukarıdaki cümlede " + stem[0].lower() + stem[1:] if stem else stem
        q["text"] = stem
        q["premise"] = sentence
        return q

    m = re.match(
        r"^(.+?\?)\s*Odak cümle:\s*(.+)$",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if m:
        q["text"] = m.group(1).replace("Aşağıdaki cümlelerden", "Yukarıdaki cümlelerden")
        q["premise"] = m.group(2).strip()
        return q

    m = re.match(
        r'^"(.+?)"\s*cümlesinde(.+)$',
        text,
        flags=re.IGNORECASE,
    )
    if m and not premise:
        q["premise"] = m.group(1).strip()
        q["text"] = "Yukarıdaki cümlede" + m.group(2).strip()
        return q

    m = re.match(
        r"^(.+?)\.\s+cümlesinde(.+)$",
        text,
        flags=re.IGNORECASE,
    )
    if m and not premise:
        q["premise"] = m.group(1).strip() + "."
        q["text"] = "Yukarıdaki cümlede" + m.group(2).strip()
        return q

    if premise and "Yukarı" not in text and "yukarı" not in text.lower():
        if parse_roman_items(premise):
            if "hangileri" in text.lower() or "hangisi" in text.lower():
                q["text"] = "Yukarıdaki " + text[0].lower() + text[1:] if text else text
            else:
                q["text"] = "Yukarıdaki bilgilere göre " + text[0].lower() + text[1:] if text else text
        else:
            prefix = premise_text_prefix(premise)
            if prefix:
                q["text"] = prefix + " " + text[0].lower() + text[1:] if text else text

    if premise and "İncelenen boşluk:" in q.get("text", ""):
        q["text"] = re.sub(r"\s*İncelenen boşluk:.*$", "", q["text"]).strip()

    return q


def split_stem(text: str, premise: str | None) -> str:
    t = (text or "").strip()
    p = (premise or "").strip()
    if not p:
        return t
    if t.startswith(p):
        rest = t[len(p) :].strip()
        return rest or "Yukarıdaki bilgilere göre doğru seçeneği işaretleyiniz."
    if len(parse_roman_items(p)) >= 2:
        for marker in (
            "Yukarıdaki cümlelerin",
            "Yukarıdaki cümlelerden",
            "Yukarıdaki eşleştirmelerden",
            "Yukarıdaki bilgilerden",
            "Yukarıdaki ifadelerden",
            "Yukarıdaki kelimelerden",
            "Yukarıdaki sözcüklerden",
        ):
            idx = t.find(marker)
            if idx >= 0:
                return t[idx:].strip()
    return t


def to_app_payload(q: dict) -> dict:
    q = normalize_question(q)
    premise = (q.get("premise") or "").strip() or None
    stem = split_stem(q.get("text") or "", premise)
    info_blocks = None
    if premise:
        roman = parse_roman_items(premise)
        if roman:
            info_blocks = [{"type": "items", "items": roman}]
        else:
            info_blocks = [{"type": "text", "content": premise}]
    return {
        "question": {
            "text": stem,
            "info": None,
            "infoItems": None,
            "infoBlocks": info_blocks,
        },
        "correct": q["correct"],
        "wrong1": q["wrong1"],
        "wrong2": q["wrong2"],
        "explanation": q["explanation"],
        "url": None,
    }


def patch_dllwrld(
    questions: list[dict],
    *,
    dllwrld_path: Path,
    topic_id: str,
    heading_id: str = "SINIF3",
    lesson_id: str = "lesson_turkce",
    lesson_id_kw: str | None = None,
    topic_name: str | None = None,
    topic_order: int | None = None,
) -> None:
    if lesson_id_kw is not None:
        lesson_id = lesson_id_kw
    with open(dllwrld_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    topics = data["championData"]["headings"][heading_id]["lessons"][lesson_id]["topics"]
    if topic_id not in topics:
        topics[topic_id] = {"active": True, "questions": {}, "questionIds": {}}
    topic_path = topics[topic_id]
    questions_out = {}
    question_ids = {}
    for q in questions:
        qid = q["id"]
        questions_out[qid] = to_app_payload(q)
        question_ids[qid] = True
    topic_path["questions"] = questions_out
    topic_path["questionIds"] = question_ids
    topic_path["active"] = True
    if topic_name:
        topic_path["name"] = topic_name
    if topic_order is not None:
        topic_path["order"] = topic_order
    with open(dllwrld_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, separators=(",", ":"))


def write_json_pack(
    questions: list[dict],
    *,
    data_path: Path,
    topic_id: str,
    title: str,
    grade: int = 3,
    heading_id: str = "SINIF3",
    lesson_id: str = "lesson_turkce",
    lesson_id_kw: str | None = None,
) -> None:
    if lesson_id_kw is not None:
        lesson_id = lesson_id_kw
    data_path.parent.mkdir(parents=True, exist_ok=True)
    out = {
        "meta": {
            "topicId": topic_id,
            "headingId": heading_id,
            "lessonId": lesson_id,
            "title": title,
            "grade": grade,
            "count": len(questions),
        },
        "questions": [{**q, "app": to_app_payload(q)} for q in questions],
    }
    with open(data_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)


def write_word_doc(
    questions: list[dict],
    *,
    word_path: Path,
    title: str,
    key_prefix: str,
    subtitle: str = "Konu: Eş anlamlı, zıt anlamlı ve sesteş (eş sesli) kelimeler",
    subject_label: str = "3. Sınıf Türkçe",
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    word_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"{subject_label} — 150 Soru (Kolay / Orta / Zor)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(subtitle)
    doc.add_paragraph("")

    level_counts = {"kolay": 0, "orta": 0, "zor": 0}
    for q in questions:
        level_counts[q["level"]] += 1
    doc.add_paragraph(
        f"Toplam: {len(questions)} soru | Kolay: {level_counts['kolay']} | "
        f"Orta: {level_counts['orta']} | Zor: {level_counts['zor']}"
    )
    doc.add_paragraph("")

    for q in questions:
        doc.add_heading(
            f"Soru {q['id'].replace(key_prefix, '')} — {q['level'].upper()}",
            level=2,
        )
        if q.get("premise"):
            p = doc.add_paragraph()
            run = p.add_run("Öncül:\n")
            run.bold = True
            roman = parse_roman_items(q["premise"])
            if roman:
                for it in roman:
                    doc.add_paragraph(f"{it['label']} {it['text']}", style="List Bullet")
            else:
                p.add_run(q["premise"])
        doc.add_paragraph(q["text"])
        doc.add_paragraph(f"A) {q['correct']}  ✓")
        doc.add_paragraph(f"B) {q['wrong1']}")
        doc.add_paragraph(f"C) {q['wrong2']}")
        exp = doc.add_paragraph()
        er = exp.add_run(f"Açıklama: {q['explanation']}")
        er.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        doc.add_paragraph("")

    doc.save(str(word_path))
