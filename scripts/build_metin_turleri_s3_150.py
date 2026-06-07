#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""3. sınıf Metin Türleri — 150 soru üret, dllwrld.json güncelle, Word dosyası oluştur."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "metin-turleri-s3-150.json"
DLLWRLD_PATH = ROOT / "dllwrld.json"
WORD_DIR = Path.home() / "Desktop" / "SORULAR WORD"
WORD_PATH = WORD_DIR / "Metin Türleri (3. Sınıf) - 150 Soru - GÜNCEL v3.docx"

TOPIC_ID = "t01"
HEADING_ID = "SINIF3"
LESSON_ID = "lesson_turkce"
KEY_PREFIX = "metin3_"


def Q(
    num: int,
    level: str,
    text: str,
    correct: str,
    wrong1: str,
    wrong2: str,
    explanation: str,
    premise: str | None = None,
) -> dict:
    return {
        "id": f"{KEY_PREFIX}{num:03d}",
        "level": level,
        "premise": premise,
        "text": text,
        "correct": correct,
        "wrong1": wrong1,
        "wrong2": wrong2,
        "explanation": explanation,
    }


def build_questions() -> list[dict]:
    import importlib.util

    bank_path = Path(__file__).resolve().parent / "metin_turleri_s3_bank.py"
    spec = importlib.util.spec_from_file_location("metin_turleri_s3_bank", bank_path)
    bank = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(bank)
    return [normalize_question(q) for q in bank.get_questions()]


PASSAGE_STEM_RE = re.compile(
    r"^(Yukarıdaki metin parçası hangi türe girer\?|"
    r"Yukarıdaki metne göre hangisi söylenebilir\?|"
    r"Yukarıdaki metne göre hangisi söylenemez\?)\s*\n+\s*(.+)$",
    re.DOTALL,
)


def strip_passage_guillemets(s: str) -> str:
    t = (s or "").strip()
    while t.startswith("«"):
        t = t[1:].strip()
    while t.endswith("»"):
        t = t[:-1].strip()
    return t.replace("«", "").replace("»", "").strip()


def normalize_question(q: dict) -> dict:
    """Örnek metni öncüle taşır, «» işaretlerini kaldırır."""
    out = dict(q)
    text = (out.get("text") or "").strip()
    premise = (out.get("premise") or "").strip() or None

    if not premise:
        m = PASSAGE_STEM_RE.match(text)
        if m:
            out["text"] = m.group(1).strip()
            premise = strip_passage_guillemets(m.group(2).strip())

    if premise:
        out["premise"] = strip_passage_guillemets(premise)
        stem = (out.get("text") or "").strip()
        if stem.startswith("Aşağıdaki metin parçası"):
            out["text"] = stem.replace("Aşağıdaki metin parçası", "Yukarıdaki metin parçası", 1)
        elif stem.startswith("Bu metne göre"):
            out["text"] = stem.replace("Bu metne göre", "Yukarıdaki metne göre", 1)

    return out


def parse_roman_items(text: str) -> list[dict]:
    items = []
    for line in (text or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^((?:I{1,3}|IV|VI{0,3}|IX|X)\.)\s*(.+)$", line)
        if m:
            items.append({"label": m.group(1), "text": m.group(2).strip()})
    return items


def split_stem(text: str, premise: str | None) -> str:
    t = (text or "").strip()
    p = (premise or "").strip()
    if not p:
        return t
    if t.startswith(p):
        rest = t[len(p) :].strip()
        return rest or "Yukarıdaki bilgilere göre doğru seçeneği işaretleyiniz."
    if len(parse_roman_items(p)) >= 2:
        for marker in (
            "Yukarıdaki eşleştirmelerden",
            "Yukarıdaki bilgilerden",
            "Yukarıdaki ifadelerden",
        ):
            idx = t.find(marker)
            if idx >= 0:
                return t[idx:].strip()
    return t


def to_app_payload(q: dict) -> dict:
    premise = (q.get("premise") or "").strip() or None
    stem = split_stem(q.get("text") or "", premise)
    info_blocks = None
    if premise:
        roman = parse_roman_items(premise)
        if roman:
            info_blocks = [{"type": "items", "items": roman}]
        else:
            info_blocks = [{"type": "text", "content": premise}]
    return {
        "question": {
            "text": stem,
            "info": None,
            "infoItems": None,
            "infoBlocks": info_blocks,
        },
        "correct": q["correct"],
        "wrong1": q["wrong1"],
        "wrong2": q["wrong2"],
        "explanation": q["explanation"],
        "url": None,
    }


def patch_dllwrld(questions: list[dict]) -> None:
    with open(DLLWRLD_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    topic_path = (
        data["championData"]["headings"][HEADING_ID]["lessons"][LESSON_ID]["topics"][TOPIC_ID]
    )
    questions_out = {}
    question_ids = {}
    for q in questions:
        qid = q["id"]
        questions_out[qid] = to_app_payload(q)
        question_ids[qid] = True
    topic_path["questions"] = questions_out
    topic_path["questionIds"] = question_ids
    topic_path["active"] = True
    with open(DLLWRLD_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, separators=(",", ":"))


def write_json(questions: list[dict]) -> None:
    DATA_PATH.parent.mkdir(parents=True, exist_ok=True)
    out = {
        "meta": {
            "topicId": TOPIC_ID,
            "headingId": HEADING_ID,
            "lessonId": LESSON_ID,
            "title": "Metin Türleri (Hikâye edici metinler, bilgilendirici metinler ve şiir)",
            "grade": 3,
            "count": len(questions),
        },
        "questions": [{**q, "app": to_app_payload(q)} for q in questions],
    }
    with open(DATA_PATH, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)


def write_word(questions: list[dict]) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = doc.add_heading(
        "Metin Türleri (Hikâye edici metinler, bilgilendirici metinler ve şiir)",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("3. Sınıf Türkçe — 150 Soru (Kolay / Orta / Zor)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Konu: Hikâye edici metinler, bilgilendirici metinler ve şiir")
    doc.add_paragraph("")

    level_counts = {"kolay": 0, "orta": 0, "zor": 0}
    for q in questions:
        level_counts[q["level"]] += 1

    doc.add_paragraph(
        f"Toplam: {len(questions)} soru | Kolay: {level_counts['kolay']} | Orta: {level_counts['orta']} | Zor: {level_counts['zor']}"
    )
    doc.add_paragraph("")

    for q in questions:
        h = doc.add_heading(f"Soru {q['id'].replace(KEY_PREFIX, '')} — {q['level'].upper()}", level=2)
        if q.get("premise"):
            p = doc.add_paragraph()
            run = p.add_run("Öncül:\n")
            run.bold = True
            roman = parse_roman_items(q["premise"])
            if roman:
                for it in roman:
                    doc.add_paragraph(f"{it['label']} {it['text']}", style="List Bullet")
            else:
                p.add_run(q["premise"])
        doc.add_paragraph(q["text"])
        doc.add_paragraph(f"A) {q['correct']}  ✓")
        doc.add_paragraph(f"B) {q['wrong1']}")
        doc.add_paragraph(f"C) {q['wrong2']}")
        exp = doc.add_paragraph()
        er = exp.add_run(f"Açıklama: {q['explanation']}")
        er.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        doc.add_paragraph("")

    doc.save(str(WORD_PATH))


def main() -> int:
    questions = build_questions()
    write_json(questions)
    patch_dllwrld(questions)
    write_word(questions)
    print(f"OK: {len(questions)} soru")
    print(f"JSON: {DATA_PATH}")
    print(f"dllwrld: {TOPIC_ID}")
    print(f"Word: {WORD_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
